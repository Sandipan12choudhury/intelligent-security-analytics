"""
==============================================================================
Export Service

Purpose
-------
Formats a single observation, together with whatever AI artifacts
have been generated for it, into a downloadable CSV, PDF, or Word
(.docx) file.

This module does NOT:

    • Call the LLM
    • Read Excel datasets directly
    • Decide which artifacts to generate

Author
------
Sandipan Choudhury
==============================================================================
"""

import csv
import io
from datetime import datetime, timezone

from docx import Document
from docx.shared import Pt, RGBColor
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle
)

from backend.models.export_models import ExportRequest

# =============================================================================
# Friendly labels for artifact types
# =============================================================================

ARTIFACT_LABELS = {

    "ROOT_CAUSE": "Root Cause",

    "RECOMMENDATION": "Recommendation",

    "BUSINESS_IMPACT": "Likely Impact",

    "EXECUTIVE_SUMMARY": "Executive Summary",

    "RISK_JUSTIFICATION": "Risk Justification",

    "PREVENTIVE_CONTROL": "Preventive Controls"

}

ARTIFACT_ORDER = [

    "BUSINESS_IMPACT",

    "RISK_JUSTIFICATION",

    "ROOT_CAUSE",

    "RECOMMENDATION",

    "PREVENTIVE_CONTROL",

    "EXECUTIVE_SUMMARY"

]


class ExportService:

    # ============================================================
    # CSV
    # ============================================================

    def to_csv(self, request: ExportRequest) -> bytes:

        buffer = io.StringIO()

        writer = csv.writer(buffer)

        writer.writerow(["Field", "Value"])

        writer.writerow(["Observation ID", request.observation_id])

        writer.writerow(["Severity", request.severity or ""])

        writer.writerow(["Domain", request.domain or ""])

        writer.writerow(["Activity", request.activity or ""])

        writer.writerow(["Application", request.application_name or ""])

        writer.writerow(["Department", request.department or ""])

        writer.writerow(["Observation", request.observation])

        writer.writerow([])

        writer.writerow(["AI Artifact", "Generated Content"])

        for artifact_type in ARTIFACT_ORDER:

            if artifact_type in request.artifacts:

                writer.writerow([

                    ARTIFACT_LABELS.get(artifact_type, artifact_type),

                    request.artifacts[artifact_type]

                ])

        return buffer.getvalue().encode("utf-8-sig")

    # ============================================================
    # PDF
    # ============================================================

    def to_pdf(self, request: ExportRequest) -> bytes:

        buffer = io.BytesIO()

        doc = SimpleDocTemplate(

            buffer,

            pagesize=A4,

            topMargin=1.5 * cm,

            bottomMargin=1.5 * cm,

            leftMargin=1.8 * cm,

            rightMargin=1.8 * cm

        )

        styles = getSampleStyleSheet()

        title_style = ParagraphStyle(

            "TitleStyle",

            parent=styles["Heading1"],

            fontSize=18,

            textColor=colors.HexColor("#0f172a"),

            spaceAfter=6

        )

        heading_style = ParagraphStyle(

            "HeadingStyle",

            parent=styles["Heading3"],

            textColor=colors.HexColor("#2563eb"),

            spaceBefore=14,

            spaceAfter=6

        )

        body_style = ParagraphStyle(

            "BodyStyle",

            parent=styles["BodyText"],

            fontSize=10.5,

            leading=15,

            textColor=colors.HexColor("#1e293b")

        )

        elements = []

        elements.append(

            Paragraph(
                "Intelligent Security Analytics Platform",
                title_style
            )

        )

        elements.append(

            Paragraph(
                f"AI Observation Report - {request.observation_id}",
                heading_style
            )

        )

        elements.append(Spacer(1, 8))

        meta_table_data = [
            ["Severity", request.severity or "-"],
            ["Domain", request.domain or "-"],
            ["Activity", request.activity or "-"],
            ["Application", request.application_name or "-"],
            ["Department", request.department or "-"],
        ]

        meta_table = Table(
            meta_table_data,
            colWidths=[4 * cm, 12 * cm]
        )

        meta_table.setStyle(TableStyle([
            ("FONTSIZE", (0, 0), (-1, -1), 9.5),
            ("TEXTCOLOR", (0, 0), (0, -1), colors.HexColor("#64748b")),
            ("TEXTCOLOR", (1, 0), (1, -1), colors.HexColor("#0f172a")),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 6),
            ("LINEBELOW", (0, 0), (-1, -1), 0.5, colors.HexColor("#e2e8f0")),
        ]))

        elements.append(meta_table)

        elements.append(Spacer(1, 14))

        elements.append(Paragraph("Observation", heading_style))

        elements.append(Paragraph(request.observation, body_style))

        for artifact_type in ARTIFACT_ORDER:

            if artifact_type in request.artifacts:

                elements.append(

                    Paragraph(

                        ARTIFACT_LABELS.get(artifact_type, artifact_type),

                        heading_style

                    )

                )

                text = request.artifacts[artifact_type].replace("\n", "<br/>")

                elements.append(Paragraph(text, body_style))

        elements.append(Spacer(1, 18))

        elements.append(

            Paragraph(

                f"Generated on "
                f"{datetime.now(timezone.utc).strftime('%d %b %Y, %H:%M')} "
                f"by Intelligent Security Analytics Platform (AI-assisted).",

                ParagraphStyle(
                    "FooterStyle",
                    parent=styles["Normal"],
                    fontSize=8,
                    textColor=colors.HexColor("#94a3b8")
                )

            )

        )

        doc.build(elements)

        return buffer.getvalue()

    # ============================================================
    # DOCX
    # ============================================================

    def to_docx(self, request: ExportRequest) -> bytes:

        document = Document()

        title = document.add_heading(
            "Intelligent Security Analytics Platform",
            level=1
        )

        title.runs[0].font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

        subtitle = document.add_heading(
            f"AI Observation Report - {request.observation_id}",
            level=2
        )

        subtitle.runs[0].font.color.rgb = RGBColor(0x25, 0x63, 0xEB)

        meta_table = document.add_table(rows=0, cols=2)

        meta_fields = [
            ("Severity", request.severity or "-"),
            ("Domain", request.domain or "-"),
            ("Activity", request.activity or "-"),
            ("Application", request.application_name or "-"),
            ("Department", request.department or "-"),
        ]

        for label, value in meta_fields:

            row = meta_table.add_row().cells

            row[0].text = label

            row[1].text = str(value)

        document.add_paragraph("")

        document.add_heading("Observation", level=2)

        document.add_paragraph(request.observation)

        for artifact_type in ARTIFACT_ORDER:

            if artifact_type in request.artifacts:

                document.add_heading(

                    ARTIFACT_LABELS.get(artifact_type, artifact_type),

                    level=2

                )

                document.add_paragraph(request.artifacts[artifact_type])

        document.add_paragraph("")

        footer = document.add_paragraph(

            f"Generated on "
            f"{datetime.now(timezone.utc).strftime('%d %b %Y, %H:%M')} "
            f"by Intelligent Security Analytics Platform (AI-assisted)."

        )

        footer.runs[0].font.size = Pt(8)

        footer.runs[0].font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

        buffer = io.BytesIO()

        document.save(buffer)

        return buffer.getvalue()
